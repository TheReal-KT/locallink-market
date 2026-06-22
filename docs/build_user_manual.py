from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "LocalLink-User-Manual.docx"
ASSET_DIR = ROOT / "docs" / "manual-assets"

COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "header_fill": "E8EEF5",
    "light_fill": "F4F6F9",
    "muted": "555555",
    "border": "D9DEE5",
}


def set_font(run, name="Calibri", size=11, bold=False, color="000000", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"], size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:color"), color)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["dark_blue"], 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)


def add_title_block(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("LocalLink Market User Manual")
    set_font(run, size=24, bold=True, color=COLORS["dark_blue"])

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run(
        "Customer and administrator guide for the LocalLink PHP/MySQL marketplace."
    )
    set_font(subtitle_run, size=11, color=COLORS["muted"])

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "This manual explains the implemented LocalLink Market website, the technology used to build it, "
        "the main functions available to customers and administrators, and the exact steps required to use each core workflow."
    )
    set_font(intro_run)


def add_section_heading(doc, text, level=1):
    doc.add_paragraph(text, style=f"Heading {level}")


def write_paragraph(doc, text, bold_label=None):
    paragraph = doc.add_paragraph()
    if bold_label:
        label_run = paragraph.add_run(f"{bold_label}: ")
        set_font(label_run, bold=True, color=COLORS["dark_blue"])
    run = paragraph.add_run(text)
    set_font(run)


def add_table(doc, headers, rows, widths, header_fill=COLORS["header_fill"]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        shade_cell(cell, header_fill)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(header)
        set_font(run, size=10.5, bold=True, color=COLORS["dark_blue"])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(value))
            set_font(run, size=10.5)

    return table


def add_screenshot(doc, filename, caption):
    image_path = ASSET_DIR / filename
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    max_width = 6.0
    max_height = 6.8
    width_in = max_width
    height_in = width_in * (height_px / width_px)
    if height_in > max_height:
        height_in = max_height
        width_in = height_in * (width_px / height_px)

    doc.add_picture(str(image_path), width=Inches(width_in), height=Inches(height_in))
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    set_font(run, size=9, color=COLORS["muted"], italic=True)


def add_operational_section(doc, title, overview, steps, screenshot_file, caption):
    add_section_heading(doc, title, level=2)
    write_paragraph(doc, overview)
    add_table(doc, ["Step", "Action"], steps, [1.15, 5.35], header_fill=COLORS["light_fill"])
    add_screenshot(doc, screenshot_file, caption)


def build_document():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_section_heading(doc, "1. System Overview", level=1)
    write_paragraph(
        doc,
        "LocalLink Market is a simplified college ecommerce platform built around three main areas: a storefront for product browsing, a customer account area for ordering and order history, and a small administrator area for basic store oversight.",
    )
    write_paragraph(
        doc,
        "The current scope intentionally focuses on the implemented customer and administrator flows. A seller toggle remains visible in navigation, but the submission-ready feature set documented here is the streamlined storefront, checkout, dashboard, and admin toolset.",
    )

    add_section_heading(doc, "2. Access and Demo Accounts", level=1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Primary local URL", "http://localhost/locallink-market/"],
            ["Customer demo login", "buyer@locallink.market / Buyer123!"],
            ["Administrator demo login", "admin@locallink.market / Admin123!"],
            ["Database setup files", "database/schema.sql and database/seed.sql"],
            ["Local deployment helper", "deploy-xampp.ps1"],
        ],
        [2.1, 4.4],
    )

    add_section_heading(doc, "3. Technical Stack", level=1)
    write_paragraph(
        doc,
        "The table below lists the technologies used in the LocalLink site and the purpose of each part of the stack.",
    )
    add_table(
        doc,
        ["Layer", "Technology", "How It Is Used In LocalLink"],
        [
            ["Backend language", "PHP", "Renders the pages, handles form submissions, manages login sessions, and processes checkout and admin actions."],
            ["Web server", "Apache via XAMPP", "Hosts the application locally under localhost for development and demonstration."],
            ["Database", "MySQL / MariaDB", "Stores users, categories, products, and orders in the locallink_market database."],
            ["Database access", "PDO", "Connects PHP to MySQL with prepared statements and exception-based error handling."],
            ["Frontend structure", "HTML5", "Defines the page layouts for the home page, product pages, login/register screens, checkout, and dashboards."],
            ["Styling", "CSS3", "Provides the LocalLink visual design through assets/css/styling.css, including layouts, cards, forms, and dashboard styling."],
            ["Client-side behaviour", "Vanilla JavaScript", "Handles the mobile navigation toggle and live checkout total updates in assets/js/app.js."],
            ["Authentication", "PHP sessions + password hashing", "Tracks signed-in users with session data and verifies passwords securely."],
            ["Data model", "users, categories, products, orders tables", "Supports the implemented marketplace flows for accounts, catalog browsing, and purchasing."],
            ["Media assets", "Local SVG/JPG assets", "Supplies product illustrations and images stored under assets/images/."],
        ],
        [1.25, 1.45, 3.8],
    )

    add_section_heading(doc, "4. System Features", level=1)
    write_paragraph(
        doc,
        "LocalLink provides separate feature sets for customers and administrators. The current version is intentionally smaller than the earlier multi-role concept so the implemented workflows remain clear and manageable.",
    )

    add_section_heading(doc, "4.1 Customer Features", level=2)
    add_table(
        doc,
        ["Feature", "Description"],
        [
            ["Homepage and category browsing", "Customers can open the homepage, review the featured products, and move into the catalog by clicking category cards or Browse Products."],
            ["Product catalog search and filters", "The products page supports keyword search, category filtering, and price/newest sorting."],
            ["Product detail page", "Each product page shows the item image, category, description, stock status, and a direct Buy now action."],
            ["Customer registration", "New users can create a customer account using name, email address, and password."],
            ["Secure login", "Returning users sign in with email and password and are redirected to the correct dashboard."],
            ["Checkout flow", "Customers choose quantity, delivery method, payment option, and optional notes before placing an order."],
            ["Order history dashboard", "Signed-in customers can review order totals, order status, and account details from the buyer dashboard."],
        ],
        [2.1, 4.4],
    )

    add_section_heading(doc, "4.2 Administrator Features", level=2)
    add_table(
        doc,
        ["Feature", "Description"],
        [
            ["Admin-only dashboard access", "Administrator accounts are protected and redirected to the admin dashboard after login."],
            ["Store statistics", "The dashboard summarises total users, products, orders, and revenue."],
            ["Latest customer orders", "Administrators can review a recent orders table showing order code, customer, quantity, status, total, and date."],
            ["Latest user signups", "The dashboard lists recently created user accounts with role and joined date."],
            ["Add product form", "Administrators can create new products by entering the product title, category, price, stock level, and short description."],
            ["Catalog management scope", "The current admin scope is intentionally focused on product creation and reporting instead of advanced moderation or seller approval workflows."],
        ],
        [2.1, 4.4],
    )

    add_section_heading(doc, "5. Operational Guide", level=1)
    write_paragraph(
        doc,
        "This section explains how to use each major workflow on the site. The screenshots were captured from the running LocalLink application hosted locally through XAMPP.",
    )

    add_operational_section(
        doc,
        "5.1 Open the Homepage and Browse the Storefront",
        "Use the homepage to understand the project scope, move into the catalog, and jump into registration or login.",
        [
            ["1", "Open the site in a browser using http://localhost/locallink-market/."],
            ["2", "Review the hero section to see the project summary and the main call-to-action buttons."],
            ["3", "Click Browse Products to open the catalog, or click Create Account if you are a new customer."],
            ["4", "Use the category cards lower on the page if you want to browse a specific product category immediately."],
        ],
        "01-home.png",
        "Figure 1. LocalLink homepage with category shortcuts, featured products, and account entry points.",
    )

    add_operational_section(
        doc,
        "5.2 Search, Filter, and Sort Products",
        "The products page is the main catalog screen for finding the correct item quickly.",
        [
            ["1", "Open Products from the top navigation bar or the homepage."],
            ["2", "Type a keyword in the Search field if you know part of the product name."],
            ["3", "Choose a category from the Category dropdown if you want to narrow the list."],
            ["4", "Use the Sort dropdown to arrange products by newest first or lowest price."],
            ["5", "Click Apply filters to refresh the list, or Clear to reset the catalog view."],
        ],
        "02-products.png",
        "Figure 2. Product catalog with keyword search, category filter, and sorting controls.",
    )

    add_operational_section(
        doc,
        "5.3 View Product Details",
        "Each product card links to a dedicated detail page where the customer can inspect the item before buying it.",
        [
            ["1", "From the homepage or products page, click on a product card or product title."],
            ["2", "Review the product image, category label, price, stock level, and description."],
            ["3", "Use Buy now to move directly into checkout for that item, or Back to products to continue browsing."],
        ],
        "03-product-detail.png",
        "Figure 3. Product detail page with pricing, stock information, and the Buy now action.",
    )

    add_operational_section(
        doc,
        "5.4 Create a Customer Account",
        "Public registration is available for customer accounts only. Administrator accounts are seeded directly through the database.",
        [
            ["1", "Open Register from the top navigation or homepage."],
            ["2", "Enter your full name, valid email address, and a password with at least eight characters."],
            ["3", "Click Create account to submit the form."],
            ["4", "After a successful registration, continue to the login page and sign in with the new customer account."],
        ],
        "04-register.png",
        "Figure 4. Customer registration form used to create a new public account.",
    )

    add_operational_section(
        doc,
        "5.5 Sign In to the Site",
        "Both customers and administrators use the same login page. The system redirects each user to the correct dashboard based on the account role.",
        [
            ["1", "Open the Login page from the top navigation."],
            ["2", "Enter your email address and password."],
            ["3", "Click Sign in."],
            ["4", "Customers are redirected to the buyer dashboard, while administrators are redirected to the admin dashboard."],
        ],
        "05-login.png",
        "Figure 5. Shared login page for both customer and administrator accounts.",
    )

    add_operational_section(
        doc,
        "5.6 Complete Checkout",
        "Checkout is only available to signed-in customers. The form captures order quantity, delivery, payment preference, and an optional note.",
        [
            ["1", "Log in as a customer and open a product detail page."],
            ["2", "Click Buy now to open checkout for that specific item."],
            ["3", "Confirm the quantity in the Order details section."],
            ["4", "Choose a delivery method: standard delivery, express delivery, or collection."],
            ["5", "Choose the payment method. If card is selected, fill in the simulated card fields."],
            ["6", "Add any order note if needed, then click Pay and place order."],
            ["7", "After the order is created, the system redirects back to the buyer dashboard where the order appears in history."],
        ],
        "07-checkout.png",
        "Figure 6. Checkout page showing quantity, delivery, payment, notes, and the live order summary.",
    )

    add_operational_section(
        doc,
        "5.7 Review the Buyer Dashboard",
        "The buyer dashboard acts as the customer account area and provides a record of previous orders.",
        [
            ["1", "Sign in with a customer account."],
            ["2", "Open My Account or allow the post-login redirect to take you to the dashboard."],
            ["3", "Review the summary cards for total orders, open orders, and total spent."],
            ["4", "Check the order history table for order number, product, quantity, status, total, and date."],
            ["5", "Use Browse products or Shop now if you want to return to the store."],
        ],
        "06-buyer-dashboard.png",
        "Figure 7. Buyer dashboard with account details, order summary cards, and order history.",
    )

    add_operational_section(
        doc,
        "5.8 Open and Use the Admin Dashboard",
        "The administrator dashboard is the control surface for store reporting and quick oversight.",
        [
            ["1", "Sign in with the administrator account."],
            ["2", "After login, the site redirects automatically to the admin dashboard."],
            ["3", "Review the summary cards for Users, Products, Orders, and Revenue."],
            ["4", "Use the Recent orders table to monitor current purchases."],
            ["5", "Use the Recent users table to review the latest signups."],
            ["6", "Use the left sidebar or Add Product button to move into product creation."],
        ],
        "08-admin-dashboard.png",
        "Figure 8. Administrator dashboard with store statistics, recent orders, and recent user signups.",
    )

    add_operational_section(
        doc,
        "5.9 Add a Product as an Administrator",
        "Product creation is handled from the administrator side in the current LocalLink scope.",
        [
            ["1", "Log in as an administrator and open Add product from the sidebar or dashboard button."],
            ["2", "Enter the product title."],
            ["3", "Choose the correct category from the dropdown list."],
            ["4", "Enter the price, stock quantity, and a short description."],
            ["5", "Click Save product to add the new listing to the database."],
            ["6", "Return to the admin dashboard or the storefront to confirm that the product has been created."],
        ],
        "09-add-product.png",
        "Figure 9. Administrator add-product form for catalog management.",
    )

    add_section_heading(doc, "6. Closing Notes", level=1)
    write_paragraph(
        doc,
        "LocalLink Market was implemented as a lightweight PHP/MySQL marketplace focused on the required customer and administrator flows. The manual therefore concentrates on the features that are active in the current codebase: browsing, registration, login, checkout, order history, admin reporting, and product creation.",
    )
    write_paragraph(
        doc,
        "If the system is run on a new machine, import database/schema.sql and database/seed.sql first, then start MySQL and Apache through XAMPP before using the demo accounts listed in this document.",
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
